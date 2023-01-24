from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('me.jpg', width=Inches(2.0))

# personal info
speak('   Hiiii, What is your name? ')
name = input('What is your name? ')
speak('Hello ' + name + 'how are you today. ' ) 
speak('Please enter your phone number and your email. ')
phone_number = input('Please enter your phone number: ')
email = input('Please enter your email: ')
document.add_paragraph(
    name + ' | ' + str(phone_number) + ' | ' + email)

# about me
document.add_heading('About me')
speak(name + " can you tell me about yourself?")
about_me = input("Tell me about yourself? ")
document.add_paragraph(about_me)

# education
document.add_heading('Education')
p = document.add_paragraph()

university = input('Enter university ')
fromDate = input('From Date ')
toDate = input('to Date ')

p.add_run(university + ' ').bold = True
p.add_run(fromDate + '-' + toDate + '\n').italic = True

speak(' Describe your experience at ' + university + ' ')
experienceDetails = input(' Describe your experience at ' + university + ' ')
p.add_run(experienceDetails)

# more experiences
while True:
    speak('Do you have any work experiences? ')
    hasMoreExp = input('Do you have any work experiences? ')
    if hasMoreExp.lower() == 'yes':
        company = input('Enter company ')
        fromDate = input('From Date ')
        toDate = input('to Date ')

        p.add_run('company' + ' ').bold = True
        p.add_run(fromDate + '-' + toDate + '\n').italic = True

        speak( 'Descibe your experience at ' + company + ' ')
        experienceDetails = input(
            'Descibe your experience at ' + company + ' ')
        p.add_run(experienceDetails)
    else:
        break

# skills
document.add_heading('Skills')
speak('Enter your skills: ')
skill = input('Enter your skills: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    moreSkills = input('Do you have anymore skills? ')
    if moreSkills.lower() == 'yes':
        skill = input('Enter your skills: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else: 
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
speak("CV Generated, Thank you.")
p.text = "CV Generated"
    
document.save('cv.docx')
